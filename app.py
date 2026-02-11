import os
import shutil
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from analyzer_logic import extract_text, call_gemini_ai  # Importing our working logic
from fastapi.responses import StreamingResponse
from docx import Document
import io
from typing import List

app = FastAPI()

# --- CORS SETUP ---
# This is the most compatible configuration for Chromium browsers (Bing/Chrome)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"]
)

# Create a temporary directory for uploaded files
UPLOAD_DIR = "temp_uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

@app.get("/")
def read_root():
    return {"status": "AI Resume Matcher API is running"}

@app.post("/analyze")
async def analyze_endpoint(
    resume_file: UploadFile = File(...),
    job_text: str = Form(...)
):
    """Endpoint to receive resume file and JD text, then return AI analysis."""
    
    file_path = os.path.join(UPLOAD_DIR, resume_file.filename)
    
    try:
        # 1. Save the uploaded file temporarily
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(resume_file.file, buffer)
        
        # 2. Extract text from the file
        resume_text = extract_text(file_path)
        
        if not resume_text:
            raise HTTPException(status_code=400, detail="Could not extract text from the resume.")

        # 3. Call the AI Logic
        analysis_result = call_gemini_ai(resume_text, job_text)
        
        # 4. Cleanup: Remove the temporary file
        os.remove(file_path)
        
        return analysis_result

    except Exception as e:
        # Cleanup if something goes wrong
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)


# to downloads results 
# @app.post("/download-docx")
# async def generate_docx(data: dict):
#     # Word Document create karein
#     doc = Document()
#     doc.add_heading('AI Resume Analysis Report', 0)

#     # Score Section
#     doc.add_heading('Job Match Score', level=1)
#     doc.add_paragraph(f"Result: {data.get('score', 'N/A')}")

#     # Matched Skills
#     doc.add_heading('Matched Skills', level=1)
#     for skill in data.get('matched', []):
#         doc.add_paragraph(skill, style='List Bullet')

#     # Missing Skills
#     doc.add_heading('Missing Skills', level=1)
#     for skill in data.get('missing', []):
#         doc.add_paragraph(skill, style='List Bullet')

#     # Suggestions
#     doc.add_heading('AI Suggestions', level=1)
#     for sug in data.get('suggestions', []):
#         doc.add_paragraph(sug, style='List Bullet')

#     # File ko memory mein save karein (RAM mein)
#     file_stream = io.BytesIO()
#     doc.save(file_stream)
#     file_stream.seek(0)

#     # File download ke liye return karein
#     return StreamingResponse(
#         file_stream, 
#         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         headers={"Content-Disposition": "attachment; filename=Analysis_Report.docx"}
#     )

@app.post("/download-docx")
async def generate_docx(data: dict):
    doc = Document()
    
    # Check karein ke Batch mode hai ya Single mode
    is_batch = data.get("is_batch", False)
    
    if is_batch:
        doc.add_heading('HR Batch Analysis Report', 0)
        candidates = data.get("candidates", [])
        
        # Table create karein (8 Columns matching your UI)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Headers set karein
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Email'
        hdr_cells[2].text = 'Experience'
        hdr_cells[3].text = 'Score'
        hdr_cells[4].text = 'Status'
        hdr_cells[5].text = 'Phone'

        # Data rows add karein
        for c in candidates:
            row_cells = table.add_row().cells
            row_cells[0].text = c.get('name', 'N/A')
            row_cells[1].text = c.get('email', 'N/A')
            row_cells[2].text = c.get('exp', 'N/A')
            row_cells[3].text = str(c.get('score', '0')) + "%"
            row_cells[4].text = c.get('status', 'N/A')
            row_cells[5].text = c.get('phone', 'N/A')
            
    else:
        # --- Purana Single Mode Logic ---
        doc.add_heading('AI Resume Analysis Report', 0)
        doc.add_heading('Candidate Match Details', level=1)
        doc.add_paragraph(f"Match Score: {data.get('score', 'N/A')}")
        
        doc.add_heading('Matched Skills', level=1)
        for skill in data.get('matched', []):
            doc.add_paragraph(skill, style='List Bullet')
        
        doc.add_heading('Missing Skills', level=1)
        for skill in data.get('missing', []):
            doc.add_paragraph(skill, style='List Bullet')

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Analysis_Report.docx"}
    )
 
@app.post("/analyze-batch")
async def analyze_batch(
    resume_files: List[UploadFile] = File(...),
    job_text: str = Form(...)
):
    """Processes multiple resumes for HR Batch Mode."""
    results = []
    
    for resume_file in resume_files:
        file_path = os.path.join(UPLOAD_DIR, resume_file.filename)
        try:
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(resume_file.file, buffer)
            
            resume_text = extract_text(file_path)
            if resume_text:
                analysis = call_gemini_ai(resume_text, job_text)
                # Client-side status logic
                analysis["status"] = "Accepted" if analysis.get("match_score", 0) >= 70 else "Rejected"
                results.append(analysis)
            
            os.remove(file_path)
        except Exception as e:
            if os.path.exists(file_path): os.remove(file_path)
            print(f"Error analyzing {resume_file.filename}: {e}")
            
    return results



# import os
# import shutil
# from fastapi import FastAPI, UploadFile, File, Form, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from analyzer_logic import extract_text, call_gemini_ai 

# app = FastAPI()

# # VERY IMPORTANT: This allows your HTML file to communicate with the Python server
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],  # Allows all origins
#     allow_credentials=True,
#     allow_methods=["*"],  # Allows POST, GET, etc.
#     allow_headers=["*"],  # Allows all headers
# )

# UPLOAD_DIR = "temp_uploads"
# os.makedirs(UPLOAD_DIR, exist_ok=True)

# @app.get("/")
# def home():
#     return {"message": "Server is up and running!"}

# @app.post("/analyze")
# async def analyze_endpoint(
#     resume_file: UploadFile = File(...),
#     job_text: str = Form(...)
# ):
#     file_path = os.path.join(UPLOAD_DIR, resume_file.filename)
#     try:
#         with open(file_path, "wb") as buffer:
#             shutil.copyfileobj(resume_file.file, buffer)
        
#         resume_text = extract_text(file_path)
#         if not resume_text:
#             raise HTTPException(status_code=400, detail="Text extraction failed.")

#         result = call_gemini_ai(resume_text, job_text)
#         os.remove(file_path) # Clean up
#         return result
#     except Exception as e:
#         if os.path.exists(file_path): os.remove(file_path)
#         raise HTTPException(status_code=500, detail=str(e))

# if __name__ == "__main__":
#     import uvicorn
#     # This forces the server to run on 127.0.0.1 port 8000
#     uvicorn.run(app, host="127.0.0.1", port=8000)